from flask import Flask, render_template, request, send_file
import pdfplumber
import pandas as pd
from docx import Document
import fitz  # PyMuPDF
import io

app = Flask(__name__, template_folder='../templates')

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        return "No file uploaded", 400
    
    file = request.files['file']
    format_type = request.form.get('format')
    pdf_bytes = file.read()

    try:
        # --- IMAGE CONVERSION ---
        if format_type == 'img':
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            page = doc.load_page(0) 
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2)) # Higher resolution
            img_data = pix.tobytes("png")
            return send_file(io.BytesIO(img_data), mimetype='image/png', as_attachment=True, download_name='page_1.png')

        # --- EXCEL CONVERSION ---
        elif format_type == 'xlsx':
            output = io.BytesIO()
            all_rows = []
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page in pdf.pages:
                    # Strategy 1: Look for explicit tables
                    table = page.extract_table()
                    if table:
                        all_rows.extend(table)
                    else:
                        # Strategy 2: Fallback to lines of text
                        text = page.extract_text()
                        if text:
                            for line in text.split('\n'):
                                all_rows.append([line])
            
            if not all_rows:
                return "Could not extract data (PDF might be an image)", 400
                
            df = pd.DataFrame(all_rows)
            df.to_excel(output, index=False, header=False)
            output.seek(0)
            return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', as_attachment=True, download_name='converted.xlsx')

        # --- WORD CONVERSION ---
        elif format_type == 'docx':
            output = io.BytesIO()
            doc = Document()
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text)
            
            if len(doc.paragraphs) == 0:
                return "Could not extract text (PDF might be an image)", 400
                
            doc.save(output)
            output.seek(0)
            return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name='converted.docx')

    except Exception as e:
        return f"Server Error: {str(e)}", 500

# For Vercel
app.debug = False
